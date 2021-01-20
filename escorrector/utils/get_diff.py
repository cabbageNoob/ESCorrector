'''
@Descripttion: 
@Author: cjh (492795090@qq.com)
@Date: Do not edit
'''
import os, sys
sys.path.insert(0, os.getcwd())
from escorrector.utils.file_utils import writejson2file, readjson

pwd_path=os.path.abspath(os.path.dirname(__file__))
TEST_SET_PATH = os.path.join(pwd_path, '../data/cn/test_set/test.txt')


'''
找出两个评测结果的不同
'''
def get_diff(file1_path, file2_path):
    f = open(os.path.join(pwd_path, 'diff.txt'), 'w', encoding='utf8')
    file1 = open(file1_path, 'r', encoding='utf8')
    file1_data = file1.readlines()
    file2 = open(file2_path, 'r', encoding='utf8')
    file2_data = file2.readlines()
    file1_list = []
    file2_list = []
    for i in range(0, len(file1_data), 5):
        temp1 = file1_data[i:i + 5]
        temp2 = file2_data[i:i + 5]
        file1_list.append(temp1)
        file2_list.append(temp2)
    for file1_line, file2_line in zip(file1_list, file2_list):
        if (file1_line[1] != file2_line[1]):
            f.write(file1_line[0])
            f.write(file1_line[1])
            f.write(file2_line[1])
            f.write(file1_line[2])
            f.write(file1_line[3])
            f.write(file1_line[4])
        else:
            continue

'''
比较两个对齐字符串的不同，用于构建测试集
'''
import numpy as np
def find_diff(a,b):
    diff_index = np.array(list(a))!=np.array(list(b))
    array_a = np.array(list(a))
    mistake = [(index,a[index],b[index]) for index, char in enumerate(diff_index) if char==True]
    return mistake

import difflib
def find_diff2(a, b):
    d=difflib.Differ()
    diff=d.compare(a.splitlines(),b.splitlines())
    print('\n'.join(list(diff)))

cnt = 1
def write_to_file(mistake, text):
    global cnt
    print(cnt)
    cnt += 1
    with open(TEST_SET_PATH, 'a', encoding='utf-8') as file:
        text_input=list(text)
        num = 1
        for location, wrong, correction in mistake:
            text_input.insert(int(location) + num, '（（' + correction + '））')
            num += 1
        file.write(''.join(text_input) + '\n')

details = []
def get_details(mistake, wrong_text, right_text):
    global details
    detail = dict()
    detail['wrong_text'] = wrong_text
    detail['right_text'] = right_text
    detail['detail'] = [(wrong_char, right_char, idx, idx + 1) for idx, wrong_char, right_char in mistake]
    details.append(detail)
'''
处理doc文档
'''
import docx, re
def doc_parse(file_path):
    file = docx.Document(file_path)
    p_three = [file.paragraphs[i: i + 3] for i in range(0, len(file.paragraphs), 3) if i + 3 <= len(file.paragraphs)]
    for p_1, p_2, p_3 in p_three:
        if p_2.text.startswith('应调整为：'):
            right_sent = re.sub(u"\\（.*?\\）", "", p_2.text[5:])
            # 等长
            if len(p_1.text) == len(right_sent):
                mistake = find_diff(p_1.text, right_sent)
                get_details(mistake, p_1.text, right_sent)
            # 不等长
            else:
                file = open(os.path.join(pwd_path, '../data/cn/test_set/un_file.txt'), 'a', encoding='utf8')
                file.writelines(p_1.text+'\n')
                file.writelines(p_2.text + '\n\n')
    writejson2file(details, os.path.join(pwd_path, '../data/cn/test_set/test.json'))
    

def main():
    paths = [r'D:\1研一上学期\nlpir实验室\工作\242-047\文件\反馈\1026-1106', r'D:\1研一上学期\nlpir实验室\工作\242-047\文件\反馈\1109-1127']
    for path in paths:
        fileList = os.listdir(path)
        for file in fileList:
            if file.startswith('~$'):
                continue
            doc_parse(path + '\\'+file)

        # print('\n')

if __name__ == '__main__':
    # file1_path = os.path.join(pwd_path, '../data/cn/sighan_eval/sighan_test/sighan_2013_eval.txt')
    # file2_path = os.path.join(pwd_path, '../data/cn/sighan_eval/sighan13_word/people14sighan_eval.txt')
    # get_diff(file1_path, file2_path)
    # a = '拖拖踏踏的背后，违建性质从经济、社会和生态环境问题逐渐转变为地方政府政治纪律和政治规矩问题。'
    # b = '拖拖沓沓的背后，违建性质从经济、社会和生态环境问题逐渐转变为地方政府政治纪律和政治规矩问题。'
    # # a = '为避免法院推诿扯皮现继续发生'
    # # b = '为避免法院推诿扯皮现象继续发生'
    # print(find_diff(a,b))
    # doc_path=os.path.join(pwd_path)
    main()
    # fileList = os.listdir(r'D:\1研一上学期\nlpir实验室\工作\242-047\文件\反馈\1109-1127')
    # print(1)